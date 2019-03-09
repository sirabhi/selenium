# -*- coding: utf-8 -*-
"""
Created on Thu Dec  6 22:30:45 2018

@author: admin
"""

import time
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from urllib import urlopen as uReq
from bs4 import BeautifulSoup as soup
from docx import Document

#from selenium.webdriver.support.ui import WebDriverWait
#from selenium.webdriver.common.by import By
#from selenium.webdriver.support import expected_conditions as EC

driver=webdriver.Firefox(executable_path="D:\\abhishek\\geckodriver.exe")
driver.maximize_window() #For maximizing window
 #gives an implicit wait for 20 seconds

driver.get("http://www.google.com")

#WebDriverWait(driver,10).until(EC.presence_of_element_located((By.ID),"gbqfq"))

driver.find_element_by_name("q").send_keys("computer graphics")
#driver.implicitly_wait(20)
#driver.implicitly_wait(20)
time.sleep(10)
driver.find_element_by_name("btnK").click()
time.sleep(10)
links=[]
"""urls=driver.find_element_by_css_selector('div.r')
k=urls.find_element_by_css_selector('a')

links=k.get_attribute('href')
print links
driver.get(links)
"""

urls=driver.find_elements_by_css_selector('div.srg a')
#urls=driver.find_element_by_tag_name('h3').findNext('a');
#k=urls.find_element_by_css_selector('a')
for l in urls:
    links.append(l.get_attribute('href'))
print links

my_url=links[0]
my_url2=links[4]
uClient=uReq(my_url)
uClient2=uReq(my_url2)
page_html=uClient.read()
page_html2=uClient2.read()

#for getting al the things on page
f=uReq(my_url)
g=soup(page_html,"html.parser")
s=g.get_text()
#//
uClient.close()
page_soup=soup(page_html,"html.parser")
page_soup2=soup(page_html2,"html.parser")
#containers=page_soup.findAll(text="introduction")
#print containers

word = page_soup.h1.get_text()
word2=page_soup2.h1.get_text()
print word2
word2=page_soup2.find('h1').findNext("p").get_text()
print word2
doc = Document()
p = doc.add_paragraph()

runner = p.add_run(word)
runner.bold = True
runner.italic = True
word=page_soup.find(text="Overview" ).findNext("p").get_text()
p=doc.add_paragraph(word)
#p.add_run(word)


goc=Document()
l=goc.add_paragraph(s)
doc.save('test.docx')
goc.save('test1.docx')

""""print(" ",page_soup.h1.get_text())
save=open("inventorji.txt",'w')
save.write(page_soup.h1.get_text()'\n\n')
save.close()
print("Overview:\n",page_soup.find(text="Overview").findNext("p").get_text())
save=open("inventorji.txt",'a')
save.write(page_soup.find(text="Overview").findNext("p").get_text())
save.close()
#driver.get(links[5])
"""
"""driver.find_element_by_name("q").send_keys(linkList[1])
time.sleep(10)
driver.find_element_by_name("btnK").click()
"""
"""
#driver.submit()
#driver.quit()
"""